#!/usr/bin/env python3
"""
Bot de Manutenção Automotiva e Agrícola - WEBHOOK MODE para Render.com
Multi-API com Fallback + Áudio + Diagnóstico
Cascata: Groq -> Gemini x2 -> DeepSeek -> Claude -> OpenAI -> Fallback
"""

import os
import sys
import json
import logging
import time
import base64
import subprocess
import threading
from http.server import HTTPServer, BaseHTTPRequestHandler
from pathlib import Path
import urllib.request
import urllib.parse
import urllib.error
import re

# --- Configuração via Variáveis de Ambiente ---
TELEGRAM_TOKEN = os.environ.get("TELEGRAM_TOKEN", "")
TELEGRAM_API = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}"
RENDER_URL = os.environ.get("RENDER_EXTERNAL_URL", "")

GEMINI_KEYS = []
gk1 = os.environ.get("GEMINI_KEY_1", "")
gk2 = os.environ.get("GEMINI_KEY_2", "")
if gk1: GEMINI_KEYS.append(gk1)
if gk2: GEMINI_KEYS.append(gk2)
GEMINI_MODEL = "gemini-2.0-flash"

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
OPENAI_BASE_URL = os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1")
OPENAI_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4.1-nano")

CLAUDE_API_KEY = os.environ.get("CLAUDE_API_KEY", "")
CLAUDE_MODEL = "claude-3-haiku-20240307"

DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_MODEL = "deepseek-chat"
DEEPSEEK_BASE_URL = "https://api.deepseek.com/v1"

GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GROQ_MODEL = os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile")
GROQ_BASE_URL = "https://api.groq.com/openai/v1"

KNOWLEDGE_DIR = Path("/tmp/knowledge_base")
KNOWLEDGE_DIR.mkdir(exist_ok=True)

# User-Agent necessário para evitar bloqueio Cloudflare (erro 1010)
USER_AGENT = "ManutBot/2.0 (Python urllib)"

ADMIN_ID = int(os.environ.get("ADMIN_ID", "912095382"))
FALLBACK_MESSAGE = "Dificuldade em processar todas as solicitações, procure o distribuidor!"

# Histórico de conversas por chat_id (memória)
_chat_history = {}  # {chat_id: [{"role": "user", "content": "..."}, {"role": "assistant", "content": "..."}]}
MAX_HISTORY = 10  # Máximo de pares de mensagens por chat (20 mensagens total)
HISTORY_TTL = 7200  # 2 horas sem atividade limpa o histórico
_chat_last_active = {}  # {chat_id: timestamp}

def add_to_history(chat_id, role, content):
    """Adiciona mensagem ao histórico do chat"""
    if chat_id not in _chat_history:
        _chat_history[chat_id] = []
    _chat_history[chat_id].append({"role": role, "content": content})
    # Manter apenas as últimas MAX_HISTORY*2 mensagens
    if len(_chat_history[chat_id]) > MAX_HISTORY * 2:
        _chat_history[chat_id] = _chat_history[chat_id][-(MAX_HISTORY * 2):]
    _chat_last_active[chat_id] = time.time()

def get_history(chat_id):
    """Retorna o histórico do chat, limpando se expirou"""
    if chat_id in _chat_last_active:
        if time.time() - _chat_last_active[chat_id] > HISTORY_TTL:
            _chat_history.pop(chat_id, None)
            _chat_last_active.pop(chat_id, None)
            return []
    return _chat_history.get(chat_id, [])

def clear_history(chat_id):
    """Limpa o histórico de um chat"""
    _chat_history.pop(chat_id, None)
    _chat_last_active.pop(chat_id, None)

# Estatísticas de uso
_stats = {
    "total_messages": 0,
    "total_audio": 0,
    "total_documents": 0,
    "users": {},  # {chat_id: {"name": "", "messages": 0, "last_active": 0}}
    "api_calls": {},  # {api_name: {"success": 0, "fail": 0}}
    "start_time": time.time(),
    "errors": [],  # últimos 50 erros
}
STATS_FILE = Path("/tmp/bot_stats.json")

def save_stats():
    try:
        STATS_FILE.write_text(json.dumps(_stats, default=str), encoding="utf-8")
    except:
        pass

def load_stats():
    global _stats
    try:
        if STATS_FILE.exists():
            loaded = json.loads(STATS_FILE.read_text(encoding="utf-8"))
            _stats.update(loaded)
    except:
        pass

def track_message(chat_id, user_name=""):
    _stats["total_messages"] += 1
    if str(chat_id) not in _stats["users"]:
        _stats["users"][str(chat_id)] = {"name": user_name, "messages": 0, "last_active": 0}
    _stats["users"][str(chat_id)]["messages"] += 1
    _stats["users"][str(chat_id)]["last_active"] = time.time()
    if user_name:
        _stats["users"][str(chat_id)]["name"] = user_name
    save_stats()

def track_api_call(api_name, success=True):
    if api_name not in _stats["api_calls"]:
        _stats["api_calls"][api_name] = {"success": 0, "fail": 0}
    if success:
        _stats["api_calls"][api_name]["success"] += 1
    else:
        _stats["api_calls"][api_name]["fail"] += 1
    save_stats()

def track_error(error_msg):
    _stats["errors"].append({"time": time.time(), "msg": error_msg[:200]})
    if len(_stats["errors"]) > 50:
        _stats["errors"] = _stats["errors"][-50:]
    save_stats()

# Senha do painel admin (variável de ambiente ou padrão)
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "manut2024")

# Cache de APIs com falha
_api_fail_cache = {}
FAIL_CACHE_TTL = 300
_notified_errors = set()

def mark_api_failed(name):
    _api_fail_cache[name] = time.time()

def is_api_available(name):
    if name not in _api_fail_cache:
        return True
    if time.time() - _api_fail_cache[name] > FAIL_CACHE_TTL:
        del _api_fail_cache[name]
        return True
    return False

# --- Logging ---
logging.basicConfig(
    format="%(asctime)s - %(levelname)s - %(message)s",
    level=logging.INFO,
    handlers=[logging.StreamHandler()],
)
logger = logging.getLogger(__name__)

# --- System Prompt ---
SYSTEM_PROMPT_BASE = (
    "Você é um técnico de manutenção especialista em equipamentos automotivos "
    "e agrícolas com acesso a busca na internet. Responda em português brasileiro.\n\n"
    "VOCÊ TEM ACESSO A INFORMAÇÕES DA INTERNET. Quando receber dados de busca técnica, "
    "USE-OS na resposta. NUNCA diga que 'não tem acesso' a informações.\n\n"
    "REGRAS OBRIGATÓRIAS:\n"
    "1. OBJETIVO e DIRETO. Sem introduções, saudações ou enrolação.\n"
    "2. Se receber INFORMAÇÕES DE BUSCA TÉCNICA abaixo, USE-AS como base da resposta. "
    "Extraia os dados relevantes e apresente de forma clara.\n"
    "3. NUNCA INVENTE números específicos (horas, torques, pressões, medidas). "
    "Só cite valores que estão EXPLICITAMENTE nas informações fornecidas.\n"
    "4. Se o dado específico NÃO está nas informações fornecidas, diga: "
    "'Não localizei este dado específico nas fontes disponíveis. "
    "Recomendo consultar o Manual de Reparação/Operador do [modelo], seção [sugestão].'\n"
    "5. NUNCA diga 'não tenho acesso' ou 'não posso acessar sites'. "
    "Você TEM acesso via busca automática.\n"
    "6. Formato: dado técnico direto → fonte → observação se necessário.\n"
    "7. Quando o usuário corrigir algo, agradeça e registre.\n"
    "8. Máximo 3-4 parágrafos.\n"
    "9. Use seu conhecimento técnico geral para contextualizar, mas não invente valores."
)

# --- Busca Web Técnica Avançada ---

# Marcas e fabricantes reconhecidos
BRANDS = [
    "john deere", "case", "case ih", "new holland", "massey ferguson", "valtra",
    "fendt", "claas", "kubota", "yanmar", "agrale", "stara", "jacto",
    "caterpillar", "cat", "komatsu", "volvo", "scania", "mercedes",
    "iveco", "ford", "volkswagen", "vw", "toyota", "mitsubishi",
    "cummins", "perkins", "mwm", "deutz", "sisu", "fpt",
    "bosch", "denso", "delphi", "rexroth", "parker", "danfoss",
    "zf", "dana", "carraro", "eaton", "allison",
    "baldan", "marchesan", "tatu", "jumil", "semeato", "sfil",
    "husqvarna", "stihl", "honda", "briggs", "kohler",
    "ls tractor", "mahindra", "landini", "deutz-fahr", "same",
]

# Termos técnicos que indicam necessidade de busca
TECH_TERMS = [
    "código de erro", "código falha", "dtc", "alarme", "erro",
    "torque", "especificação", "regulagem", "calibração", "ajuste",
    "intervalo", "manutenção", "troca de óleo", "filtro", "lubrificação",
    "pressão", "vazão", "rpm", "horímetro", "temperatura",
    "esquema elétrico", "diagrama", "fusível", "relé", "sensor",
    "peça", "part number", "referência", "número",
    "manual", "procedimento", "reparação", "desmontagem", "montagem",
    "válvula", "bomba", "injetor", "turbo", "embreagem", "freio",
    "hidráulico", "transmissão", "diferencial", "eixo",
    "motor", "cabeçote", "pistão", "biela", "virabrequim",
    "colhedora", "trator", "pulverizador", "plantadeira",
]

def extract_brand_model(text):
    """Extrai marca e modelo da mensagem do usuário"""
    text_lower = text.lower()
    found_brand = None
    found_model = None
    
    for brand in BRANDS:
        if brand in text_lower:
            found_brand = brand
            break
    
    # Tentar extrair modelo (números e letras após a marca)
    if found_brand:
        # Procurar padrões como "CH 570", "7200J", "MF 4275", "CR 9080"
        patterns = [
            r'(?:' + re.escape(found_brand) + r')\s*([a-zA-Z]*\s*\d+[a-zA-Z]*(?:\s*\d+)*)',
            r'([A-Z]{1,4}\s*\d{2,5}[A-Z]?(?:\s*\d{4})?)',  # CH 570 2025, 7200J
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                found_model = match.group(1).strip()
                break
    
    return found_brand, found_model

def do_web_search(search_terms, timeout=8):
    """Executa busca web - tenta DuckDuckGo e Google como fallback"""
    results = _search_duckduckgo(search_terms, timeout)
    if not results:
        results = _search_google(search_terms, timeout)
    return results

def _search_duckduckgo(search_terms, timeout=8):
    """Busca via DuckDuckGo HTML"""
    results = []
    try:
        encoded = urllib.parse.quote(search_terms)
        url = f"https://html.duckduckgo.com/html/?q={encoded}"
        
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml",
            "Accept-Language": "pt-BR,pt;q=0.9,en;q=0.8",
        })
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            html = resp.read().decode("utf-8", errors="ignore")
        
        snippets = re.findall(r'<a class="result__snippet"[^>]*>(.*?)</a>', html, re.DOTALL)
        titles = re.findall(r'<a class="result__a"[^>]*>(.*?)</a>', html, re.DOTALL)
        urls = re.findall(r'<a class="result__url"[^>]*href="([^"]+)"', html)
        
        for i, (title, snippet) in enumerate(zip(titles[:5], snippets[:5])):
            clean_title = re.sub(r'<[^>]+>', '', title).strip()
            clean_snippet = re.sub(r'<[^>]+>', '', snippet).strip()
            if clean_title and clean_snippet and len(clean_snippet) > 20:
                source = urls[i] if i < len(urls) else ""
                results.append({
                    "title": clean_title,
                    "snippet": clean_snippet,
                    "url": source,
                })
    except Exception as e:
        logger.warning(f"DuckDuckGo falhou: {e}")
    return results

def _search_google(search_terms, timeout=8):
    """Busca via Google (scraping básico como fallback)"""
    results = []
    try:
        encoded = urllib.parse.quote(search_terms)
        url = f"https://www.google.com/search?q={encoded}&hl=pt-BR&num=5"
        
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml",
            "Accept-Language": "pt-BR,pt;q=0.9,en;q=0.8",
        })
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            html = resp.read().decode("utf-8", errors="ignore")
        
        # Extrair blocos de resultado do Google
        blocks = re.findall(r'<div class="[^"]*"[^>]*>.*?</div>', html[:50000], re.DOTALL)
        
        # Tentar extrair texto útil
        text_content = re.sub(r'<[^>]+>', ' ', html[:30000])
        text_content = re.sub(r'\s+', ' ', text_content)
        
        # Extrair frases que parecem resultados (entre pontos, com mais de 50 chars)
        sentences = re.findall(r'([A-ZÀ-Ü][^.!?]{50,200}[.!?])', text_content)
        for s in sentences[:5]:
            s = s.strip()
            if len(s) > 50 and not any(x in s.lower() for x in ['google', 'cookie', 'privacy', 'javascript']):
                results.append({
                    "title": "Google",
                    "snippet": s,
                    "url": "",
                })
    except Exception as e:
        logger.warning(f"Google falhou: {e}")
    return results

def fetch_page_content(url, max_chars=3000):
    """Tenta acessar uma página web e extrair conteúdo textual relevante"""
    try:
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        })
        with urllib.request.urlopen(req, timeout=6) as resp:
            content_type = resp.headers.get('Content-Type', '')
            if 'text/html' not in content_type:
                return None
            html = resp.read(50000).decode("utf-8", errors="ignore")
        
        # Remover scripts, styles, nav
        html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
        html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL)
        html = re.sub(r'<nav[^>]*>.*?</nav>', '', html, flags=re.DOTALL)
        html = re.sub(r'<header[^>]*>.*?</header>', '', html, flags=re.DOTALL)
        html = re.sub(r'<footer[^>]*>.*?</footer>', '', html, flags=re.DOTALL)
        
        # Extrair texto
        text = re.sub(r'<[^>]+>', ' ', html)
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Filtrar conteúdo muito curto
        if len(text) < 100:
            return None
        
        return text[:max_chars]
    except:
        return None

def search_technical_info(query):
    """Busca avançada de informações técnicas - múltiplas fontes e estratégias"""
    brand, model = extract_brand_model(query)
    all_results = []
    
    # Montar buscas específicas baseadas na marca/modelo
    search_queries = []
    
    if brand and model:
        # Buscas muito específicas para o modelo
        search_queries.append(f"{brand} {model} especificações técnicas motor")
        search_queries.append(f"{brand} {model} manual reparação")
        search_queries.append(f"{brand} {model} service manual specifications")
        search_queries.append(f"{brand} {model} ficha técnica")
    elif brand:
        search_queries.append(f"{brand} {query}")
        search_queries.append(f"{brand} {query} manual técnico especificações")
        search_queries.append(f"{brand} {query} repair service")
    else:
        search_queries.append(f"{query} especificações técnicas")
        search_queries.append(f"{query} manual reparação")
    
    # Executar buscas (até 4 queries)
    seen_snippets = set()
    for sq in search_queries[:4]:
        results = do_web_search(sq)
        for r in results:
            key = r["snippet"][:50]
            if key not in seen_snippets:
                seen_snippets.add(key)
                all_results.append(r)
    
    if not all_results:
        return None
    
    # Tentar acessar até 2 páginas relevantes para obter mais detalhes
    page_contents = []
    accessed = 0
    for r in all_results[:4]:
        if accessed >= 2:
            break
        url = r.get("url", "")
        if url and not any(x in url for x in [".pdf", "youtube", "facebook", "instagram", "twitter"]):
            # Resolver URL do DuckDuckGo redirect
            if "duckduckgo.com" in url:
                parsed = urllib.parse.urlparse(url)
                params = urllib.parse.parse_qs(parsed.query)
                url = params.get("uddg", [url])[0]
            
            content = fetch_page_content(url, max_chars=2500)
            if content and len(content) > 200:
                page_contents.append(f"Fonte: {r['title']}\n{content}")
                accessed += 1
    
    # Montar resultado final
    output_parts = []
    
    # Snippets dos resultados de busca
    output_parts.append("RESULTADOS DE BUSCA (use estes dados na resposta):")
    for r in all_results[:6]:
        output_parts.append(f"• {r['title']}: {r['snippet']}")
    
    # Conteúdo das páginas acessadas
    if page_contents:
        output_parts.append("\nCONTEÚDO DETALHADO DAS FONTES (dados confiáveis):")
        for pc in page_contents:
            output_parts.append(pc)
    
    result = "\n".join(output_parts)
    logger.info(f"Busca técnica: {len(all_results)} resultados, {len(page_contents)} páginas acessadas")
    return result

def detect_equipment_query(text):
    """Detecta se a mensagem menciona equipamento/marca específica que vale buscar"""
    text_lower = text.lower()
    
    has_brand = any(b in text_lower for b in BRANDS)
    has_tech = any(t in text_lower for t in TECH_TERMS)
    
    # Buscar se tem marca OU se tem termo técnico específico
    return has_brand or has_tech

# --- Telegram API ---
def telegram_request(method, data=None, timeout=60):
    url = f"{TELEGRAM_API}/{method}"
    try:
        if data:
            data_encoded = urllib.parse.urlencode(data).encode("utf-8")
        else:
            data_encoded = None
        req = urllib.request.Request(url, data=data_encoded)
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="ignore")
        logger.error(f"HTTP Error {e.code} em {method}: {body[:200]}")
        return None
    except Exception as e:
        logger.error(f"Erro em {method}: {e}")
        return None

def send_message(chat_id, text):
    if len(text) > 4000:
        for i in range(0, len(text), 4000):
            telegram_request("sendMessage", {"chat_id": chat_id, "text": text[i:i+4000]})
    else:
        result = telegram_request("sendMessage", {"chat_id": chat_id, "text": text})
        if result and result.get("ok"):
            logger.info(f"Mensagem enviada para {chat_id}")
        else:
            logger.error(f"Falha ao enviar para {chat_id}: {result}")

def send_typing(chat_id):
    telegram_request("sendChatAction", {"chat_id": chat_id, "action": "typing"})

def download_file(file_id):
    result = telegram_request("getFile", {"file_id": file_id})
    if result and result.get("ok"):
        file_path = result["result"]["file_path"]
        url = f"https://api.telegram.org/file/bot{TELEGRAM_TOKEN}/{file_path}"
        local_path = f"/tmp/{os.path.basename(file_path)}"
        urllib.request.urlretrieve(url, local_path)
        return local_path
    return None

def notify_admin_error(api_name, error_msg):
    """Notifica o admin sobre erro de API (apenas primeira vez)"""
    key = f"{api_name}_{error_msg[:50]}"
    if key not in _notified_errors:
        _notified_errors.add(key)
        try:
            send_message(ADMIN_ID, f"⚠️ API {api_name} falhou:\n{error_msg[:300]}")
        except:
            pass

# --- APIs de IA ---
def call_gemini(prompt, api_key):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={api_key}"
    payload = json.dumps({
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"maxOutputTokens": 2000, "temperature": 0.7},
    }).encode("utf-8")
    req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json", "User-Agent": USER_AGENT})
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            result = json.loads(resp.read().decode("utf-8"))
            candidates = result.get("candidates", [])
            if candidates:
                parts = candidates[0].get("content", {}).get("parts", [])
                if parts:
                    text = parts[0].get("text", "")
                    if text:
                        return text
            return None
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="ignore")
        error_msg = f"HTTP {e.code}: {body[:150]}"
        logger.warning(f"Gemini ...{api_key[-6:]} falhou: {error_msg}")
        notify_admin_error(f"Gemini ...{api_key[-6:]}", error_msg)
        return None
    except Exception as e:
        logger.warning(f"Gemini ...{api_key[-6:]} erro: {e}")
        notify_admin_error(f"Gemini ...{api_key[-6:]}", str(e))
        return None

def call_openai_compatible(messages_list, api_key, base_url, model, name="API"):
    """Chama API compatível com OpenAI. messages_list é lista de dicts com role/content."""
    if not api_key:
        return None
    url = f"{base_url}/chat/completions"
    payload = json.dumps({
        "model": model,
        "messages": messages_list,
        "max_tokens": 1500,
        "temperature": 0.7,
    }).encode("utf-8")
    req = urllib.request.Request(url, data=payload, headers={
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
        "User-Agent": USER_AGENT,
    })
    try:
        with urllib.request.urlopen(req, timeout=45) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            choices = data.get('choices', [])
            if choices:
                return choices[0].get('message', {}).get('content', '')
            return None
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="ignore")
        error_msg = f"HTTP {e.code}: {body[:200]}"
        logger.warning(f"{name} falhou: {error_msg}")
        notify_admin_error(name, error_msg)
        return None
    except Exception as e:
        error_msg = str(e)
        logger.warning(f"{name} erro: {error_msg}")
        notify_admin_error(name, error_msg)
        return None

def call_claude(prompt):
    if not CLAUDE_API_KEY:
        return None
    url = "https://api.anthropic.com/v1/messages"
    payload = json.dumps({
        "model": CLAUDE_MODEL,
        "max_tokens": 1500,
        "messages": [{"role": "user", "content": prompt}],
    }).encode("utf-8")
    req = urllib.request.Request(url, data=payload, headers={
        "Content-Type": "application/json",
        "x-api-key": CLAUDE_API_KEY,
        "anthropic-version": "2023-06-01",
        "User-Agent": USER_AGENT,
    })
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            result = json.loads(resp.read().decode("utf-8"))
            content = result.get("content", [])
            if content:
                return content[0].get("text", "")
            return None
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="ignore")
        error_msg = f"HTTP {e.code}: {body[:150]}"
        logger.warning(f"Claude falhou: {error_msg}")
        notify_admin_error("Claude", error_msg)
        return None
    except Exception as e:
        logger.warning(f"Claude erro: {e}")
        notify_admin_error("Claude", str(e))
        return None

def call_ai_with_fallback(messages_list, gemini_prompt):
    """Tenta APIs em cascata. messages_list para APIs OpenAI-compat, gemini_prompt para Gemini."""
    # 1. Groq (gratuito e rápido)
    if GROQ_API_KEY and is_api_available("groq"):
        logger.info("Tentando Groq...")
        res = call_openai_compatible(messages_list, GROQ_API_KEY, GROQ_BASE_URL, GROQ_MODEL, "Groq")
        if res:
            logger.info(f"Groq respondeu ({len(res)} chars)")
            track_api_call("Groq", True)
            return res
        track_api_call("Groq", False)
        mark_api_failed("groq")

    # 2. Gemini (não suporta histórico multi-turn facilmente, usa prompt concatenado)
    for i, key in enumerate(GEMINI_KEYS):
        name = f"gemini_{i+1}"
        if not is_api_available(name):
            continue
        logger.info(f"Tentando Gemini key {i+1}...")
        res = call_gemini(gemini_prompt, key)
        if res:
            logger.info(f"Gemini respondeu ({len(res)} chars)")
            track_api_call(f"Gemini_{i+1}", True)
            return res
        track_api_call(f"Gemini_{i+1}", False)
        mark_api_failed(name)

    # 3. DeepSeek
    if DEEPSEEK_API_KEY and is_api_available("deepseek"):
        logger.info("Tentando DeepSeek...")
        res = call_openai_compatible(messages_list, DEEPSEEK_API_KEY, DEEPSEEK_BASE_URL, DEEPSEEK_MODEL, "DeepSeek")
        if res:
            logger.info(f"DeepSeek respondeu ({len(res)} chars)")
            track_api_call("DeepSeek", True)
            return res
        track_api_call("DeepSeek", False)
        mark_api_failed("deepseek")

    # 4. Claude
    if CLAUDE_API_KEY and is_api_available("claude"):
        logger.info("Tentando Claude...")
        res = call_claude(gemini_prompt)
        if res:
            logger.info(f"Claude respondeu ({len(res)} chars)")
            track_api_call("Claude", True)
            return res
        track_api_call("Claude", False)
        mark_api_failed("claude")

    # 5. OpenAI
    if OPENAI_API_KEY and is_api_available("openai"):
        logger.info("Tentando OpenAI...")
        res = call_openai_compatible(messages_list, OPENAI_API_KEY, OPENAI_BASE_URL, OPENAI_MODEL, "OpenAI")
        if res:
            logger.info(f"OpenAI respondeu ({len(res)} chars)")
            track_api_call("OpenAI", True)
            return res
        track_api_call("OpenAI", False)
        mark_api_failed("openai")

    logger.warning("Todas as APIs falharam!")
    track_error("Todas as APIs falharam")
    return None

# --- Transcrição de Áudio ---
def transcribe_audio(audio_path):
    wav_path = audio_path + ".wav"
    try:
        result = subprocess.run(
            ["ffmpeg", "-i", audio_path, "-ar", "16000", "-ac", "1", "-y", wav_path],
            capture_output=True, timeout=30
        )
        if result.returncode != 0:
            logger.error(f"ffmpeg erro: {result.stderr.decode()[:200]}")
            return None
    except Exception as e:
        logger.error(f"ffmpeg exception: {e}")
        return None

    # Tenta Groq Whisper primeiro (rápido e gratuito)
    if GROQ_API_KEY:
        text = transcribe_with_groq_whisper(wav_path)
        if text:
            return text

    # Tenta Gemini multimodal
    for i, key in enumerate(GEMINI_KEYS):
        name = f"gemini_stt_{i+1}"
        if not is_api_available(name):
            continue
        text = transcribe_with_gemini(wav_path, key)
        if text:
            return text
        mark_api_failed(name)

    # Tenta OpenAI Whisper
    if OPENAI_API_KEY:
        text = transcribe_with_openai_whisper(wav_path, OPENAI_API_KEY, OPENAI_BASE_URL)
        if text:
            return text

    return None

def transcribe_with_groq_whisper(wav_path):
    """Transcreve áudio usando Groq Whisper API"""
    if not GROQ_API_KEY:
        return None
    import http.client
    try:
        boundary = "----FormBoundary7MA4YWxkGroq"
        with open(wav_path, "rb") as f:
            file_data = f.read()
        body = b""
        body += f"--{boundary}\r\n".encode()
        body += b'Content-Disposition: form-data; name="file"; filename="audio.wav"\r\n'
        body += b"Content-Type: audio/wav\r\n\r\n"
        body += file_data
        body += b"\r\n"
        body += f"--{boundary}\r\n".encode()
        body += b'Content-Disposition: form-data; name="model"\r\n\r\n'
        body += b"whisper-large-v3\r\n"
        body += f"--{boundary}\r\n".encode()
        body += b'Content-Disposition: form-data; name="language"\r\n\r\n'
        body += b"pt\r\n"
        body += f"--{boundary}--\r\n".encode()
        conn = http.client.HTTPSConnection("api.groq.com", timeout=30)
        conn.request("POST", "/openai/v1/audio/transcriptions", body=body, headers={
            "Authorization": f"Bearer {GROQ_API_KEY}",
            "Content-Type": f"multipart/form-data; boundary={boundary}",
        })
        resp = conn.getresponse()
        data = json.loads(resp.read().decode("utf-8"))
        conn.close()
        if resp.status == 200:
            return data.get("text", "")
        logger.warning(f"Groq Whisper falhou (HTTP {resp.status}): {json.dumps(data)[:150]}")
        return None
    except Exception as e:
        logger.warning(f"Groq Whisper erro: {e}")
        return None

def transcribe_with_gemini(wav_path, api_key):
    try:
        with open(wav_path, "rb") as f:
            audio_data = base64.b64encode(f.read()).decode("utf-8")
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={api_key}"
        payload = json.dumps({
            "contents": [{"parts": [
                {"inline_data": {"mime_type": "audio/wav", "data": audio_data}},
                {"text": "Transcreva este áudio em português brasileiro. Retorne APENAS o texto transcrito, sem explicações."}
            ]}],
            "generationConfig": {"maxOutputTokens": 500, "temperature": 0.1},
        }).encode("utf-8")
        req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json", "User-Agent": USER_AGENT})
        with urllib.request.urlopen(req, timeout=30) as resp:
            result = json.loads(resp.read().decode("utf-8"))
            candidates = result.get("candidates", [])
            if candidates:
                parts = candidates[0].get("content", {}).get("parts", [])
                if parts:
                    return parts[0].get("text", "").strip()
        return None
    except Exception as e:
        logger.warning(f"Gemini STT erro: {e}")
        return None

def transcribe_with_openai_whisper(wav_path, api_key, base_url):
    if not api_key:
        return None
    import http.client
    host = base_url.replace("https://", "").replace("http://", "")
    if "/" in host:
        path_prefix = "/" + "/".join(host.split("/")[1:])
        host = host.split("/")[0]
    else:
        path_prefix = ""
    boundary = "----FormBoundary7MA4YWxk"
    with open(wav_path, "rb") as f:
        file_data = f.read()
    body = b""
    body += f"--{boundary}\r\n".encode()
    body += b'Content-Disposition: form-data; name="file"; filename="audio.wav"\r\n'
    body += b"Content-Type: audio/wav\r\n\r\n"
    body += file_data
    body += b"\r\n"
    body += f"--{boundary}\r\n".encode()
    body += b'Content-Disposition: form-data; name="model"\r\n\r\n'
    body += b"whisper-1\r\n"
    body += f"--{boundary}\r\n".encode()
    body += b'Content-Disposition: form-data; name="language"\r\n\r\n'
    body += b"pt\r\n"
    body += f"--{boundary}--\r\n".encode()
    try:
        conn = http.client.HTTPSConnection(host, timeout=30)
        conn.request("POST", f"{path_prefix}/audio/transcriptions", body=body, headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": f"multipart/form-data; boundary={boundary}",
        })
        resp = conn.getresponse()
        data = json.loads(resp.read().decode("utf-8"))
        conn.close()
        if resp.status == 200:
            return data.get("text", "")
        return None
    except Exception as e:
        logger.warning(f"Whisper erro: {e}")
        return None

# --- Banco de Conhecimento ---
def load_knowledge_base():
    texts = []
    for f in KNOWLEDGE_DIR.iterdir():
        if f.suffix == ".txt":
            try:
                texts.append(f.read_text(encoding="utf-8", errors="ignore"))
            except:
                pass
    return "\n\n---\n\n".join(texts) if texts else ""

def build_system_prompt():
    """Constrói o system prompt com base de conhecimento"""
    knowledge = load_knowledge_base()
    prompt = SYSTEM_PROMPT_BASE
    if knowledge:
        prompt += (
            "\n\n=== INFORMAÇÕES DOS MANUAIS CARREGADOS (PRIORIDADE MÁXIMA) ===\n"
            + knowledge[:4000]
        )
    return prompt

def build_messages_with_history(chat_id, user_text, web_info=None):
    """Constrói lista de mensagens com histórico para APIs OpenAI-compatible"""
    system_prompt = build_system_prompt()
    messages = [{"role": "system", "content": system_prompt}]
    
    # Adiciona histórico anterior
    history = get_history(chat_id)
    for msg in history:
        messages.append(msg)
    
    # Adiciona mensagem atual COM as informações de busca embutidas
    if web_info:
        enriched_text = (
            f"{user_text}\n\n"
            f"[INFORMAÇÕES TÉCNICAS ENCONTRADAS - USE ESTES DADOS NA SUA RESPOSTA]:\n"
            f"{web_info[:4000]}"
        )
        messages.append({"role": "user", "content": enriched_text})
    else:
        messages.append({"role": "user", "content": user_text})
    return messages

def build_gemini_prompt_with_history(chat_id, user_text, web_info=None):
    """Constrói prompt concatenado com histórico para Gemini"""
    system_prompt = build_system_prompt()
    prompt = system_prompt
    
    # Adiciona histórico como contexto
    history = get_history(chat_id)
    if history:
        prompt += "\n\n=== HISTÓRICO DA CONVERSA ===\n"
        for msg in history:
            role_label = "Usuário" if msg["role"] == "user" else "Assistente"
            prompt += f"{role_label}: {msg['content'][:500]}\n\n"
    
    # Adiciona a pergunta com dados de busca
    if web_info:
        prompt += (
            f"\nPergunta do usuário: {user_text}\n\n"
            f"[INFORMAÇÕES TÉCNICAS ENCONTRADAS - USE ESTES DADOS NA SUA RESPOSTA]:\n"
            f"{web_info[:4000]}"
        )
    else:
        prompt += f"\nPergunta do usuário: {user_text}"
    return prompt

# --- Handlers ---
def handle_start(chat_id):
    clear_history(chat_id)
    send_message(chat_id,
        "🔧 Especialista em Manutenção Online!\n\n"
        "Sou um assistente de IA especializado em manutenção e reparação de "
        "equipamentos automotivos e de agricultura.\n\n"
        "📋 O que posso fazer:\n"
        "• Responder dúvidas sobre manutenção preventiva e corretiva\n"
        "• Ajudar a diagnosticar problemas em equipamentos\n"
        "• Orientar sobre procedimentos de reparo\n"
        "• Lembrar do contexto da conversa (memória)\n\n"
        "📄 Aceito documentos (PDF, Word, TXT) com manuais técnicos\n"
        "🎤 Aceito mensagens de voz e áudio!\n\n"
        "📌 Comandos:\n"
        "/limpar - Limpar histórico e começar novo assunto\n"
        "/status - Ver status das APIs (admin)\n"
        "/diag - Diagnosticar APIs (admin)\n\n"
        "Como posso ajudar no diagnóstico hoje?"
    )

def handle_status(chat_id):
    """Mostra status das APIs configuradas"""
    if chat_id != ADMIN_ID:
        send_message(chat_id, "⚠️ Comando disponível apenas para o administrador.")
        return
    
    def mask_key(key):
        if not key:
            return "❌ NÃO CONFIGURADA"
        return f"✅ Configurada ({key[:8]}...{key[-4:]})"
    
    status = (
        "📊 STATUS DAS APIs:\n\n"
        f"🔹 GROQ_API_KEY: {mask_key(GROQ_API_KEY)}\n"
        f"   Modelo: {GROQ_MODEL}\n"
        f"   URL: {GROQ_BASE_URL}\n\n"
        f"🔹 GEMINI_KEY_1: {mask_key(gk1)}\n"
        f"🔹 GEMINI_KEY_2: {mask_key(gk2)}\n"
        f"   Modelo: {GEMINI_MODEL}\n\n"
        f"🔹 DEEPSEEK_API_KEY: {mask_key(DEEPSEEK_API_KEY)}\n"
        f"   Modelo: {DEEPSEEK_MODEL}\n\n"
        f"🔹 CLAUDE_API_KEY: {mask_key(CLAUDE_API_KEY)}\n"
        f"   Modelo: {CLAUDE_MODEL}\n\n"
        f"🔹 OPENAI_API_KEY: {mask_key(OPENAI_API_KEY)}\n"
        f"   Modelo: {OPENAI_MODEL}\n"
        f"   URL: {OPENAI_BASE_URL}\n\n"
        f"🔹 TELEGRAM_TOKEN: {mask_key(TELEGRAM_TOKEN)}\n"
        f"🔹 RENDER_URL: {RENDER_URL or '❌ NÃO CONFIGURADA'}\n"
        f"🔹 ADMIN_ID: {ADMIN_ID}\n\n"
        f"📁 Manuais no banco: {sum(1 for f in KNOWLEDGE_DIR.iterdir() if f.suffix == '.txt')}\n"
        f"🚫 APIs em cache de falha: {list(_api_fail_cache.keys()) if _api_fail_cache else 'nenhuma'}"
    )
    send_message(chat_id, status)

def handle_diag(chat_id):
    """Testa cada API individualmente"""
    if chat_id != ADMIN_ID:
        send_message(chat_id, "⚠️ Comando disponível apenas para o administrador.")
        return
    
    send_message(chat_id, "🔍 Iniciando diagnóstico de APIs... (pode levar até 1 minuto)")
    send_typing(chat_id)
    
    results = []
    test_prompt = "Diga apenas: OK FUNCIONANDO"
    
    # Teste Groq
    if GROQ_API_KEY:
        try:
            url = f"{GROQ_BASE_URL}/chat/completions"
            payload = json.dumps({
                "model": GROQ_MODEL,
                "messages": [{"role": "user", "content": test_prompt}],
                "max_tokens": 20,
            }).encode("utf-8")
            req = urllib.request.Request(url, data=payload, headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "User-Agent": USER_AGENT,
            })
            with urllib.request.urlopen(req, timeout=15) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                reply = data.get('choices', [{}])[0].get('message', {}).get('content', '')
                results.append(f"✅ Groq: OK - \"{reply[:50]}\"")
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="ignore")
            results.append(f"❌ Groq: HTTP {e.code} - {body[:100]}")
        except Exception as e:
            results.append(f"❌ Groq: {str(e)[:100]}")
    else:
        results.append("⚪ Groq: CHAVE NÃO CONFIGURADA")
    
    # Teste Gemini
    for i, key in enumerate(GEMINI_KEYS):
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={key}"
            payload = json.dumps({
                "contents": [{"parts": [{"text": test_prompt}]}],
                "generationConfig": {"maxOutputTokens": 20},
            }).encode("utf-8")
            req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json", "User-Agent": USER_AGENT})
            with urllib.request.urlopen(req, timeout=15) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                candidates = data.get("candidates", [])
                if candidates:
                    text = candidates[0].get("content", {}).get("parts", [{}])[0].get("text", "")
                    results.append(f"✅ Gemini key {i+1}: OK - \"{text[:50]}\"")
                else:
                    results.append(f"❌ Gemini key {i+1}: Sem candidates na resposta")
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="ignore")
            results.append(f"❌ Gemini key {i+1}: HTTP {e.code} - {body[:100]}")
        except Exception as e:
            results.append(f"❌ Gemini key {i+1}: {str(e)[:100]}")
    if not GEMINI_KEYS:
        results.append("⚪ Gemini: NENHUMA CHAVE CONFIGURADA")
    
    # Teste DeepSeek
    if DEEPSEEK_API_KEY:
        try:
            url = f"{DEEPSEEK_BASE_URL}/chat/completions"
            payload = json.dumps({
                "model": DEEPSEEK_MODEL,
                "messages": [{"role": "user", "content": test_prompt}],
                "max_tokens": 20,
            }).encode("utf-8")
            req = urllib.request.Request(url, data=payload, headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                "User-Agent": USER_AGENT,
            })
            with urllib.request.urlopen(req, timeout=15) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                reply = data.get('choices', [{}])[0].get('message', {}).get('content', '')
                results.append(f"✅ DeepSeek: OK - \"{reply[:50]}\"")
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="ignore")
            results.append(f"❌ DeepSeek: HTTP {e.code} - {body[:100]}")
        except Exception as e:
            results.append(f"❌ DeepSeek: {str(e)[:100]}")
    else:
        results.append("⚪ DeepSeek: CHAVE NÃO CONFIGURADA")
    
    # Teste Claude
    if CLAUDE_API_KEY:
        try:
            url = "https://api.anthropic.com/v1/messages"
            payload = json.dumps({
                "model": CLAUDE_MODEL,
                "max_tokens": 20,
                "messages": [{"role": "user", "content": test_prompt}],
            }).encode("utf-8")
            req = urllib.request.Request(url, data=payload, headers={
                "Content-Type": "application/json",
                "x-api-key": CLAUDE_API_KEY,
                "anthropic-version": "2023-06-01",
                "User-Agent": USER_AGENT,
            })
            with urllib.request.urlopen(req, timeout=15) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                content = data.get("content", [])
                if content:
                    results.append(f"✅ Claude: OK - \"{content[0].get('text', '')[:50]}\"")
                else:
                    results.append(f"❌ Claude: Resposta vazia")
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="ignore")
            results.append(f"❌ Claude: HTTP {e.code} - {body[:100]}")
        except Exception as e:
            results.append(f"❌ Claude: {str(e)[:100]}")
    else:
        results.append("⚪ Claude: CHAVE NÃO CONFIGURADA")
    
    # Teste OpenAI
    if OPENAI_API_KEY:
        try:
            url = f"{OPENAI_BASE_URL}/chat/completions"
            payload = json.dumps({
                "model": OPENAI_MODEL,
                "messages": [{"role": "user", "content": test_prompt}],
                "max_tokens": 20,
            }).encode("utf-8")
            req = urllib.request.Request(url, data=payload, headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {OPENAI_API_KEY}",
                "User-Agent": USER_AGENT,
            })
            with urllib.request.urlopen(req, timeout=15) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                reply = data.get('choices', [{}])[0].get('message', {}).get('content', '')
                results.append(f"✅ OpenAI: OK - \"{reply[:50]}\"")

        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="ignore")
            results.append(f"❌ OpenAI: HTTP {e.code} - {body[:100]}")
        except Exception as e:
            results.append(f"❌ OpenAI: {str(e)[:100]}")
    else:
        results.append("⚪ OpenAI: CHAVE NÃO CONFIGURADA")
    
    # Limpar cache de falhas
    _api_fail_cache.clear()
    _notified_errors.clear()
    
    report = "🔍 RESULTADO DO DIAGNÓSTICO:\n\n" + "\n\n".join(results)
    report += "\n\n🔄 Cache de falhas limpo."
    send_message(chat_id, report)

def handle_text(chat_id, text):
    # Comando para limpar histórico
    if text.strip().lower() in ["/limpar", "/novo", "/reset"]:
        clear_history(chat_id)
        send_message(chat_id, "🔄 Histórico limpo! Pode começar um novo assunto.")
        return
    
    logger.info(f"Texto de {chat_id}: {text[:80]}")
    send_typing(chat_id)
    
    # Busca web técnica - SEMPRE tenta buscar para perguntas técnicas
    web_info = None
    if detect_equipment_query(text):
        logger.info(f"Busca técnica ativada para: {text[:50]}")
        web_info = search_technical_info(text)
        if web_info:
            logger.info(f"Busca retornou {len(web_info)} chars")
        else:
            logger.warning(f"Busca não retornou resultados para: {text[:50]}")
    else:
        # Mesmo sem marca/termo específico, tenta buscar se parece pergunta técnica
        if len(text) > 15 and "?" in text:
            web_info = search_technical_info(text)
            if web_info:
                logger.info(f"Busca geral retornou {len(web_info)} chars")
    
    # Construir mensagens com histórico + info da web
    messages = build_messages_with_history(chat_id, text, web_info)
    gemini_prompt = build_gemini_prompt_with_history(chat_id, text, web_info)
    
    reply = call_ai_with_fallback(messages, gemini_prompt)
    
    if reply:
        # Salvar no histórico
        add_to_history(chat_id, "user", text)
        add_to_history(chat_id, "assistant", reply)
        send_message(chat_id, reply)
    else:
        send_message(chat_id, FALLBACK_MESSAGE)

def handle_voice(chat_id, voice_or_audio):
    file_id = voice_or_audio.get("file_id")
    logger.info(f"Áudio de {chat_id}")
    send_typing(chat_id)
    send_message(chat_id, "🎤 Processando seu áudio...")
    try:
        local_path = download_file(file_id)
        if not local_path:
            send_message(chat_id, "⚠️ Não consegui baixar o áudio.")
            return
        transcribed = transcribe_audio(local_path)
        if not transcribed:
            send_message(chat_id, "⚠️ Não consegui transcrever o áudio. Tente enviar como texto.")
            return
        send_message(chat_id, f"🎤 Entendi: \"{transcribed}\"\n\nProcessando resposta...")
        send_typing(chat_id)
        
        # Busca web técnica se detectar equipamento
        web_info = None
        if detect_equipment_query(transcribed):
            web_info = search_technical_info(transcribed)
        
        # Construir mensagens com histórico + info da web
        messages = build_messages_with_history(chat_id, transcribed, web_info)
        gemini_prompt = build_gemini_prompt_with_history(chat_id, transcribed, web_info)
        
        reply = call_ai_with_fallback(messages, gemini_prompt)
        
        if reply:
            add_to_history(chat_id, "user", transcribed)
            add_to_history(chat_id, "assistant", reply)
            send_message(chat_id, reply)
        else:
            send_message(chat_id, FALLBACK_MESSAGE)
    except Exception as e:
        logger.error(f"Erro áudio: {e}", exc_info=True)
        send_message(chat_id, "Desculpe, erro ao processar o áudio.")

def handle_document(chat_id, document):
    if chat_id != ADMIN_ID:
        send_message(chat_id, "⚠️ Apenas o administrador pode enviar documentos para o banco de conhecimento.")
        return
    file_name = document.get("file_name", "doc")
    send_message(chat_id, f"📥 Processando manual: {file_name}")
    try:
        local_path = download_file(document.get("file_id"))
        if not local_path:
            send_message(chat_id, "❌ Não consegui baixar o arquivo.")
            return
        text = ""
        if file_name.lower().endswith(".pdf"):
            from PyPDF2 import PdfReader
            reader = PdfReader(local_path)
            text = "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
        elif file_name.lower().endswith(".docx"):
            from docx import Document
            doc = Document(local_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        elif file_name.lower().endswith(".txt"):
            with open(local_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            send_message(chat_id, "⚠️ Formato não suportado. Envie PDF, Word (.docx) ou TXT.")
            return
        if text.strip():
            safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in file_name)
            save_path = KNOWLEDGE_DIR / f"{safe_name}.txt"
            save_path.write_text(text, encoding="utf-8")
            send_message(chat_id, f"✅ Manual '{file_name}' indexado com sucesso!\n📊 {len(text)} caracteres.")
        else:
            send_message(chat_id, "❌ Não foi possível extrair texto do arquivo.")
    except Exception as e:
        logger.error(f"Erro doc: {e}", exc_info=True)
        send_message(chat_id, "❌ Erro técnico ao processar documento.")

# --- Processar Update ---
def process_update(update):
    try:
        msg = update.get("message", {})
        chat_id = msg.get("chat", {}).get("id")
        if not chat_id:
            return
        
        # Tracking de usuário
        user = msg.get("from", {})
        user_name = user.get("first_name", "") + " " + user.get("last_name", "")
        user_name = user_name.strip() or str(chat_id)
        track_message(chat_id, user_name)
        
        text = msg.get("text", "")
        if text.startswith("/start"):
            handle_start(chat_id)
        elif text.startswith("/status"):
            handle_status(chat_id)
        elif text.startswith("/diag"):
            handle_diag(chat_id)
        elif msg.get("voice"):
            _stats["total_audio"] += 1
            handle_voice(chat_id, msg["voice"])
        elif msg.get("audio"):
            _stats["total_audio"] += 1
            handle_voice(chat_id, msg["audio"])
        elif msg.get("document"):
            _stats["total_documents"] += 1
            handle_document(chat_id, msg["document"])
        elif text:
            handle_text(chat_id, text)
    except Exception as e:
        logger.error(f"Erro update: {e}", exc_info=True)
        track_error(str(e))

# --- Webhook HTTP Server ---
def check_admin_auth(handler):
    """Verifica autenticação básica para o painel admin"""
    auth = handler.headers.get('Authorization', '')
    if auth.startswith('Basic '):
        try:
            decoded = base64.b64decode(auth[6:]).decode('utf-8')
            user, pwd = decoded.split(':', 1)
            if user == 'admin' and pwd == ADMIN_PASSWORD:
                return True
        except:
            pass
    return False

def send_auth_required(handler):
    handler.send_response(401)
    handler.send_header('WWW-Authenticate', 'Basic realm="Admin Panel"')
    handler.send_header('Content-Type', 'text/html')
    handler.end_headers()
    handler.wfile.write(b'<h1>Acesso negado</h1><p>Credenciais inv\xc3\xa1lidas.</p>')

def get_admin_html():
    uptime = time.time() - _stats.get("start_time", time.time())
    hours = int(uptime // 3600)
    minutes = int((uptime % 3600) // 60)
    
    # Usuários ativos (últimas 24h)
    now = time.time()
    active_24h = sum(1 for u in _stats["users"].values() if now - u.get("last_active", 0) < 86400)
    
    # Manuais
    manuals = []
    for f in KNOWLEDGE_DIR.iterdir():
        if f.suffix == ".txt":
            size = f.stat().st_size
            manuals.append({"name": f.stem, "size": size})
    
    # API stats
    api_rows = ""
    for api_name, counts in _stats.get("api_calls", {}).items():
        total = counts["success"] + counts["fail"]
        rate = (counts["success"] / total * 100) if total > 0 else 0
        api_rows += f'<tr><td>{api_name}</td><td>{counts["success"]}</td><td>{counts["fail"]}</td><td>{rate:.0f}%</td></tr>'
    
    # Usuários
    user_rows = ""
    sorted_users = sorted(_stats.get("users", {}).items(), key=lambda x: x[1].get("messages", 0), reverse=True)
    for uid, udata in sorted_users[:20]:
        last = time.strftime("%d/%m %H:%M", time.localtime(udata.get("last_active", 0))) if udata.get("last_active") else "-"
        user_rows += f'<tr><td>{udata.get("name", uid)}</td><td>{uid}</td><td>{udata.get("messages", 0)}</td><td>{last}</td></tr>'
    
    # Manuais
    manual_rows = ""
    for m in manuals:
        size_kb = m["size"] / 1024
        manual_rows += f'<tr><td>{m["name"]}</td><td>{size_kb:.1f} KB</td><td><a href="/admin/manual/delete?name={urllib.parse.quote(m["name"])}" onclick="return confirm(\'Remover este manual?\')">🗑️ Remover</a></td></tr>'
    
    # Erros recentes
    error_rows = ""
    for err in reversed(_stats.get("errors", [])[-10:]):
        t_str = time.strftime("%d/%m %H:%M:%S", time.localtime(err.get("time", 0)))
        error_rows += f'<tr><td>{t_str}</td><td>{err.get("msg", "")}</td></tr>'
    
    # APIs configuradas
    apis_status = ""
    apis_config = [
        ("Groq", bool(GROQ_API_KEY)),
        ("Gemini 1", bool(gk1)),
        ("Gemini 2", bool(gk2)),
        ("DeepSeek", bool(DEEPSEEK_API_KEY)),
        ("Claude", bool(CLAUDE_API_KEY)),
        ("OpenAI", bool(OPENAI_API_KEY)),
    ]
    for name, configured in apis_config:
        icon = "🟢" if configured else "🔴"
        apis_status += f'<span style="margin-right:15px">{icon} {name}</span>'

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>🔧 Bot Manutenção - Painel Admin</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; background: #1a1a2e; color: #eee; padding: 20px; }}
        .container {{ max-width: 1200px; margin: 0 auto; }}
        h1 {{ color: #00d4aa; margin-bottom: 5px; font-size: 1.8em; }}
        h2 {{ color: #00d4aa; margin: 25px 0 10px; font-size: 1.3em; border-bottom: 1px solid #333; padding-bottom: 5px; }}
        .subtitle {{ color: #888; margin-bottom: 20px; }}
        .cards {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; margin: 20px 0; }}
        .card {{ background: #16213e; border-radius: 10px; padding: 20px; text-align: center; }}
        .card .number {{ font-size: 2.5em; font-weight: bold; color: #00d4aa; }}
        .card .label {{ color: #888; margin-top: 5px; }}
        table {{ width: 100%; border-collapse: collapse; margin: 10px 0; background: #16213e; border-radius: 8px; overflow: hidden; }}
        th, td {{ padding: 10px 15px; text-align: left; border-bottom: 1px solid #2a2a4a; }}
        th {{ background: #0f3460; color: #00d4aa; }}
        tr:hover {{ background: #1a1a3e; }}
        .apis-bar {{ background: #16213e; padding: 15px; border-radius: 8px; margin: 10px 0; }}
        a {{ color: #ff6b6b; text-decoration: none; }}
        a:hover {{ text-decoration: underline; }}
        .upload-form {{ background: #16213e; padding: 20px; border-radius: 8px; margin: 10px 0; }}
        input[type="file"] {{ margin: 10px 0; }}
        button {{ background: #00d4aa; color: #1a1a2e; border: none; padding: 10px 20px; border-radius: 5px; cursor: pointer; font-weight: bold; }}
        button:hover {{ background: #00b894; }}
        .refresh {{ float: right; font-size: 0.9em; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>🔧 Bot Especialista em Manutenção</h1>
        <p class="subtitle">Painel Administrativo | Uptime: {hours}h {minutes}min</p>
        
        <div class="cards">
            <div class="card">
                <div class="number">{_stats.get('total_messages', 0)}</div>
                <div class="label">Mensagens</div>
            </div>
            <div class="card">
                <div class="number">{len(_stats.get('users', {}))}</div>
                <div class="label">Usuários Total</div>
            </div>
            <div class="card">
                <div class="number">{active_24h}</div>
                <div class="label">Ativos (24h)</div>
            </div>
            <div class="card">
                <div class="number">{_stats.get('total_audio', 0)}</div>
                <div class="label">Áudios</div>
            </div>
            <div class="card">
                <div class="number">{len(manuals)}</div>
                <div class="label">Manuais</div>
            </div>
            <div class="card">
                <div class="number">{_stats.get('total_documents', 0)}</div>
                <div class="label">Docs Recebidos</div>
            </div>
        </div>

        <h2>🔌 APIs Configuradas</h2>
        <div class="apis-bar">{apis_status}</div>

        <h2>📊 Uso das APIs</h2>
        <table>
            <tr><th>API</th><th>Sucesso</th><th>Falha</th><th>Taxa</th></tr>
            {api_rows if api_rows else '<tr><td colspan="4">Nenhuma chamada registrada ainda</td></tr>'}
        </table>

        <h2>👥 Usuários</h2>
        <table>
            <tr><th>Nome</th><th>ID</th><th>Mensagens</th><th>Último Acesso</th></tr>
            {user_rows if user_rows else '<tr><td colspan="4">Nenhum usuário ainda</td></tr>'}
        </table>

        <h2>📚 Manuais no Banco de Conhecimento</h2>
        <table>
            <tr><th>Nome</th><th>Tamanho</th><th>Ação</th></tr>
            {manual_rows if manual_rows else '<tr><td colspan="3">Nenhum manual carregado</td></tr>'}
        </table>
        
        <div class="upload-form">
            <h3 style="color:#00d4aa;margin-bottom:10px">📤 Upload de Manual</h3>
            <form action="/admin/upload" method="POST" enctype="multipart/form-data">
                <input type="file" name="file" accept=".pdf,.docx,.txt" required>
                <button type="submit">Enviar Manual</button>
            </form>
        </div>

        <h2>⚠️ Erros Recentes</h2>
        <table>
            <tr><th>Data/Hora</th><th>Erro</th></tr>
            {error_rows if error_rows else '<tr><td colspan="2">Nenhum erro registrado</td></tr>'}
        </table>
        
        <p style="margin-top:30px;color:#555;text-align:center">Bot Manutenção v2.3 | <a href="/admin" style="color:#00d4aa">🔄 Atualizar</a></p>
    </div>
</body>
</html>"""
    return html

class WebhookHandler(BaseHTTPRequestHandler):
    def do_GET(self):
        path = self.path.split('?')[0]
        
        if path == '/admin':
            if not check_admin_auth(self):
                send_auth_required(self)
                return
            html = get_admin_html()
            self.send_response(200)
            self.send_header('Content-Type', 'text/html; charset=utf-8')
            self.end_headers()
            self.wfile.write(html.encode('utf-8'))
            return
        
        if path == '/admin/manual/delete':
            if not check_admin_auth(self):
                send_auth_required(self)
                return
            query = urllib.parse.urlparse(self.path).query
            params = urllib.parse.parse_qs(query)
            name = params.get('name', [''])[0]
            if name:
                file_path = KNOWLEDGE_DIR / f"{name}.txt"
                if file_path.exists():
                    file_path.unlink()
            self.send_response(302)
            self.send_header('Location', '/admin')
            self.end_headers()
            return
        
        if path == '/admin/stats.json':
            if not check_admin_auth(self):
                send_auth_required(self)
                return
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps(_stats, default=str).encode('utf-8'))
            return
        
        # Página padrão
        self.send_response(200)
        self.send_header('Content-Type', 'text/plain')
        self.end_headers()
        self.wfile.write(b'Bot de Manutencao - Online')

    def do_POST(self):
        path = self.path.split('?')[0]
        
        # Upload de manual via painel admin
        if path == '/admin/upload':
            if not check_admin_auth(self):
                send_auth_required(self)
                return
            try:
                content_type = self.headers.get('Content-Type', '')
                if 'multipart/form-data' in content_type:
                    # Parse multipart
                    boundary = content_type.split('boundary=')[1].strip()
                    content_length = int(self.headers.get('Content-Length', 0))
                    body = self.rfile.read(content_length)
                    
                    # Extrair arquivo do multipart
                    parts = body.split(f'--{boundary}'.encode())
                    for part in parts:
                        if b'filename="' in part:
                            # Extrair nome do arquivo
                            header_end = part.find(b'\r\n\r\n')
                            header = part[:header_end].decode('utf-8', errors='ignore')
                            file_data = part[header_end+4:]
                            if file_data.endswith(b'\r\n'):
                                file_data = file_data[:-2]
                            
                            # Extrair filename
                            fname_start = header.find('filename="') + 10
                            fname_end = header.find('"', fname_start)
                            filename = header[fname_start:fname_end]
                            
                            if filename:
                                # Salvar temporariamente
                                tmp_path = f"/tmp/upload_{filename}"
                                with open(tmp_path, 'wb') as f:
                                    f.write(file_data)
                                
                                # Extrair texto
                                text = ""
                                if filename.lower().endswith('.pdf'):
                                    from PyPDF2 import PdfReader
                                    reader = PdfReader(tmp_path)
                                    text = "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
                                elif filename.lower().endswith('.docx'):
                                    from docx import Document as DocxDoc
                                    doc = DocxDoc(tmp_path)
                                    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                                elif filename.lower().endswith('.txt'):
                                    with open(tmp_path, 'r', encoding='utf-8', errors='ignore') as f:
                                        text = f.read()
                                
                                if text.strip():
                                    safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in filename)
                                    save_path = KNOWLEDGE_DIR / f"{safe_name}.txt"
                                    save_path.write_text(text, encoding='utf-8')
                                
                                # Limpar tmp
                                try:
                                    os.remove(tmp_path)
                                except:
                                    pass
                            break
            except Exception as e:
                logger.error(f"Erro upload admin: {e}")
            
            self.send_response(302)
            self.send_header('Location', '/admin')
            self.end_headers()
            return
        
        # Webhook do Telegram
        if path == '/webhook':
            try:
                content_length = int(self.headers.get('Content-Length', 0))
                body = self.rfile.read(content_length)
                update = json.loads(body.decode('utf-8'))
                logger.info(f"Webhook update {update.get('update_id', '?')}")

                # Responde 200 imediatamente
                self.send_response(200)
                self.send_header('Content-Type', 'text/plain')
                self.end_headers()
                self.wfile.write(b'OK')

                # Processa em thread separada para não bloquear
                t = threading.Thread(target=process_update, args=(update,))
                t.daemon = True
                t.start()
            except Exception as e:
                logger.error(f"Erro webhook: {e}")
                self.send_response(500)
                self.end_headers()
            return
        
        # POST genérico
        try:
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length)
            update = json.loads(body.decode('utf-8'))
            self.send_response(200)
            self.send_header('Content-Type', 'text/plain')
            self.end_headers()
            self.wfile.write(b'OK')
            t = threading.Thread(target=process_update, args=(update,))
            t.daemon = True
            t.start()
        except Exception as e:
            logger.error(f"Erro POST: {e}")
            self.send_response(500)
            self.end_headers()

    def log_message(self, format, *args):
        pass  # Silencia logs HTTP padrão

def setup_webhook():
    if not RENDER_URL:
        logger.error("RENDER_EXTERNAL_URL não configurado! Adicione esta variável no Render.")
        return False

    webhook_url = f"{RENDER_URL}/webhook"
    logger.info(f"Configurando webhook: {webhook_url}")

    # Remove webhook antigo
    telegram_request("deleteWebhook", {"drop_pending_updates": "false"})
    time.sleep(1)

    # Configura novo webhook
    result = telegram_request("setWebhook", {
        "url": webhook_url,
        "allowed_updates": json.dumps(["message"]),
    })

    if result and result.get("ok"):
        logger.info("Webhook configurado com sucesso!")
        return True
    else:
        logger.error(f"Falha webhook: {result}")
        return False

def main():
    if not TELEGRAM_TOKEN:
        logger.error("TELEGRAM_TOKEN não configurado!")
        sys.exit(1)

    # Carregar estatísticas salvas
    load_stats()
    
    logger.info("=" * 50)
    logger.info("BOT MANUTENÇÃO - WEBHOOK MODE v2.3")
    logger.info(f"Render URL: {RENDER_URL}")
    logger.info(f"Admin ID: {ADMIN_ID}")
    logger.info(f"Groq: {'sim' if GROQ_API_KEY else 'não'}")
    logger.info(f"Gemini keys: {len(GEMINI_KEYS)}")
    logger.info(f"DeepSeek: {'sim' if DEEPSEEK_API_KEY else 'não'}")
    logger.info(f"Claude: {'sim' if CLAUDE_API_KEY else 'não'}")
    logger.info(f"OpenAI: {'sim' if OPENAI_API_KEY else 'não'}")
    logger.info(f"Admin Panel: {RENDER_URL}/admin (user: admin)")
    logger.info("=" * 50)

    me = telegram_request("getMe")
    if not me or not me.get("ok"):
        logger.error(f"Token inválido: {me}")
        sys.exit(1)
    logger.info(f"Bot: @{me['result']['username']}")

    if not setup_webhook():
        logger.error("Falha ao configurar webhook! Verifique RENDER_EXTERNAL_URL.")
        sys.exit(1)

    port = int(os.environ.get('PORT', 10000))
    server = HTTPServer(('0.0.0.0', port), WebhookHandler)
    logger.info(f"Servidor webhook na porta {port} - PRONTO!")

    try:
        server.serve_forever()
    except KeyboardInterrupt:
        server.shutdown()

if __name__ == "__main__":
    main()
